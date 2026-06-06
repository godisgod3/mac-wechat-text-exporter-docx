#!/usr/bin/env python3
from __future__ import annotations

import argparse
import re
import shutil
import subprocess
import tempfile
from collections import Counter
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


MSG_RE = re.compile(r"^\[(\d{4}-\d{2}-\d{2} \d{2}:\d{2})\]\s(.*)$")
BAD_XML_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f]")


def clean(value: str) -> str:
    return BAD_XML_RE.sub("", value)


def set_run(run, font: str = "Microsoft YaHei", size: float = 9.0, color: str = "000000", bold=None):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def shade(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def run_wx_export(chat: str, output_txt: Path, limit: int, since: str | None, until: str | None):
    if shutil.which("wx"):
        cmd = ["wx"]
    else:
        cmd = ["npx", "-y", "@jackwener/wx-cli"]

    cmd += ["export", chat, "--format", "txt", "-n", str(limit), "-o", str(output_txt)]
    if since:
        cmd += ["--since", since]
    if until:
        cmd += ["--until", until]
    subprocess.run(cmd, check=True)


def parse_messages(source: Path):
    current = None
    with source.open("r", encoding="utf-8", errors="replace") as f:
        for raw in f:
            line = raw.rstrip("\n")
            if line.startswith("==="):
                continue
            match = MSG_RE.match(line)
            if match:
                if current:
                    yield current
                body = match.group(2)
                sender = "系统/无发送者"
                content = body
                split = re.match(r"^(.{1,120}?):\s?(.*)$", body)
                if split:
                    sender = split.group(1).strip()
                    content = split.group(2)
                current = {
                    "time": match.group(1),
                    "sender": clean(sender),
                    "content": clean(content),
                }
                continue
            if current is not None:
                current["content"] += "\n" + clean(line)
        if current:
            yield current


def configure(doc: Document):
    section = doc.sections[0]
    section.page_width = Cm(21.59)
    section.page_height = Cm(27.94)
    section.top_margin = Cm(1.45)
    section.bottom_margin = Cm(1.45)
    section.left_margin = Cm(1.65)
    section.right_margin = Cm(1.65)
    section.header_distance = Cm(0.9)
    section.footer_distance = Cm(0.9)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(9)
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1.05

    for name, size, color in [
        ("Heading 1", 16, "1F4D78"),
        ("Heading 2", 12, "2E74B5"),
        ("Heading 3", 10, "1F4D78"),
    ]:
        style = doc.styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(10 if name == "Heading 1" else 6)
        style.paragraph_format.space_after = Pt(4)


def add_meta_table(doc: Document, rows: list[tuple[str, str]]):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].width = Cm(3.6)
        cells[1].width = Cm(12.0)
        cells[0].text = label
        cells[1].text = value
        shade(cells[0], "F2F4F7")
        for cell in cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(1)
                for r in p.runs:
                    set_run(r, size=9)


def add_message(doc: Document, msg: dict[str, str]):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1.5)
    p.paragraph_format.line_spacing = 1.03

    r = p.add_run(f"[{msg['time']}] ")
    set_run(r, "Consolas", 8.2, "666666")

    r = p.add_run(f"{msg['sender']}: ")
    set_run(r, "Microsoft YaHei", 8.7, "1F4D78", True)

    content = msg["content"].strip() or "[空消息]"
    r = p.add_run(content)
    set_run(r, "Microsoft YaHei", 8.7, "000000")


def build_docx(source_txt: Path, chat: str, output_docx: Path, account_label: str):
    messages = list(parse_messages(source_txt))
    if not messages:
        raise SystemExit(f"No messages parsed from {source_txt}")

    days = Counter(m["time"][:10] for m in messages)
    senders = Counter(m["sender"] for m in messages)

    doc = Document()
    configure(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    r = title.add_run(f"{chat} 完整文字聊天记录")
    set_run(r, size=18, color="0B2545", bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(12)
    r = subtitle.add_run(f"导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}    消息数：{len(messages)}")
    set_run(r, size=9.5, color="555555")

    add_meta_table(
        doc,
        [
            ("聊天对象", chat),
            ("消息总数", f"{len(messages)} 条"),
            ("时间范围", f"{messages[0]['time']} 至 {messages[-1]['time']}"),
            ("来源账号", account_label),
            ("说明", "仅包含文字化聊天记录；不包含图片、视频、语音、文件附件本体。"),
        ],
    )

    doc.add_heading("按日期概览", level=1)
    add_meta_table(doc, [(day, f"{num} 条") for day, num in sorted(days.items())])

    doc.add_heading("发言人概览（前 30）", level=1)
    add_meta_table(doc, [(sender, f"{num} 条") for sender, num in senders.most_common(30)])

    doc.add_page_break()
    doc.add_heading("完整聊天记录", level=1)

    current_day = None
    for msg in messages:
        day = msg["time"][:10]
        if day != current_day:
            current_day = day
            doc.add_heading(day, level=2)
        add_message(doc, msg)

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_docx)


def main():
    parser = argparse.ArgumentParser(description="Export a local Mac WeChat chat/group timeline to DOCX.")
    parser.add_argument("--chat", required=True, help="Chat or group name as recognized by wx-cli.")
    parser.add_argument("--output", required=True, type=Path, help="Output DOCX path.")
    parser.add_argument("--source-txt", type=Path, help="Existing wx-cli txt export. If omitted, wx-cli export is run first.")
    parser.add_argument("--limit", type=int, default=999999, help="Maximum messages to ask wx-cli to export.")
    parser.add_argument("--since", help="Start date, YYYY-MM-DD.")
    parser.add_argument("--until", help="End date, YYYY-MM-DD.")
    parser.add_argument("--account-label", default="本机 wx-cli 当前配置", help="Non-sensitive label written into the DOCX metadata table.")
    parser.add_argument("--keep-txt", type=Path, help="Optional path to keep the intermediate TXT export.")
    args = parser.parse_args()

    if args.source_txt:
        source_txt = args.source_txt
        build_docx(source_txt, args.chat, args.output, args.account_label)
        print(args.output)
        return

    with tempfile.TemporaryDirectory(prefix="wechat_export_") as tmp:
        tmp_txt = Path(tmp) / "chat.txt"
        run_wx_export(args.chat, tmp_txt, args.limit, args.since, args.until)
        if args.keep_txt:
            args.keep_txt.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(tmp_txt, args.keep_txt)
        build_docx(tmp_txt, args.chat, args.output, args.account_label)
    print(args.output)


if __name__ == "__main__":
    main()

